import os
import re
import logging
from typing import List, Dict, Optional

logger = logging.getLogger(__name__)

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
    logger.warning("python-docx o'rnatilmagan — DOCX qo'llab-quvvatlanmaydi")

try:
    import pdfplumber
except ImportError:
    pdfplumber = None
    logger.warning("pdfplumber o'rnatilmagan — PDF qo'llab-quvvatlanmaydi")


# ─────────────────────────────── Extractors ──────────────────────────────────

def extract_text_from_docx(path: str) -> str:
    if DocxDocument is None:
        raise RuntimeError("python-docx o'rnatilmagan")
    try:
        doc = DocxDocument(path)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        parts.append(text)
        return '\n'.join(parts)
    except Exception as e:
        raise RuntimeError(f"DOCX o'qishda xato: {e}")


def extract_text_from_pdf(path: str) -> str:
    if pdfplumber is None:
        raise RuntimeError("pdfplumber o'rnatilmagan")
    try:
        texts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    texts.append(page_text)
        return '\n'.join(texts)
    except Exception as e:
        raise RuntimeError(f"PDF o'qishda xato: {e}")


# ────────────────────────────────── Parser ───────────────────────────────────

# Savol boshlanishini aniqlash: 1. 1) 1- 1: kabi
Q_RE = re.compile(r'^(\d{1,4})\s*[.)\-:]\s+')
# Variant aniqlash: A. A) A- A: kabi
OPT_RE = re.compile(r'^([A-Da-d])\s*[.)\-:]\s*')
# Javob aniqlash turli formatlarda
ANS_RE = re.compile(
    r'(?:Javob|To\'g\'ri\s+javob|Answer|Correct\s+answer|Correct|To\'g\'ri)[:\s]+([A-Da-d])',
    re.IGNORECASE
)


def clean_option_text(text: str) -> str:
    """Variant matnini tozalaydi"""
    return text.strip().rstrip('.')


def parse_text_to_questions(text: str) -> List[Dict]:
    """
    Matndan test savollarini ajratib oladi.
    Qo'llab-quvvatlangan formatlar:
      1. Savol matni
      A) Variant A
      B) Variant B
      C) Variant C
      D) Variant D
      Javob: B
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    questions = []
    current: Optional[Dict] = None

    for line in lines:
        # Javob satri
        ans_match = ANS_RE.search(line)
        if ans_match and current is not None:
            current['correct'] = ans_match.group(1).upper()
            continue

        # Savol boshi
        q_match = Q_RE.match(line)
        if q_match:
            if current is not None:
                questions.append(current)
            q_text = Q_RE.sub('', line).strip()
            current = {
                'question': q_text,
                'a': None, 'b': None, 'c': None, 'd': None,
                'correct': None
            }
            continue

        if current is None:
            continue

        # Variant satri
        opt_match = OPT_RE.match(line)
        if opt_match:
            key = opt_match.group(1).lower()
            opt_text = clean_option_text(OPT_RE.sub('', line))
            if key in ('a', 'b', 'c', 'd'):
                current[key] = opt_text
            continue

        # Ko'p qatorli savol matni (variant yo'q bo'lsa)
        if current is not None and not any([current['a'], current['b'], current['c'], current['d']]):
            # Faqat savol matniga qo'shamiz
            if len(line) > 0:
                current['question'] += ' ' + line

    # Oxirgi savolni qo'shish
    if current is not None:
        questions.append(current)

    # Filterlash: kamida bitta variant bo'lishi kerak
    valid = []
    for q in questions:
        if not q.get('question', '').strip():
            continue
        opts = [q.get(k) for k in ('a', 'b', 'c', 'd')]
        if any(o for o in opts if o and o.strip()):
            valid.append(q)
        else:
            logger.debug(f"Savol o'tkazib yuborildi (variantlar yo'q): {q['question'][:50]}")

    logger.info(f"Jami {len(valid)} ta savol ajratildi (jami {len(questions)} ta topildi)")
    return valid


# ─────────────────────────────── Public API ──────────────────────────────────

def parse_file(path: str) -> List[Dict]:
    """
    Faylni o'qib savollarni qaytaradi.
    Qo'llab-quvvatlangan formatlar: .pdf, .docx
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"Fayl topilmadi: {path}")

    _, ext = os.path.splitext(path.lower())

    if ext == '.docx':
        text = extract_text_from_docx(path)
    elif ext == '.pdf':
        text = extract_text_from_pdf(path)
    else:
        raise RuntimeError(f"Qo'llab-quvvatlanmaydigan format: {ext}. Faqat PDF va DOCX.")

    if not text or not text.strip():
        raise RuntimeError("Fayldan matn ajratib olinmadi (bo'sh yoki o'qib bo'lmaydigan fayl)")

    questions = parse_text_to_questions(text)
    return questions
